import os
import docx

path = "/mnt/c/Users/razer/Downloads/Telegram Desktop/Уведомление  АБДИЕВ 2026.docx"

if not os.path.exists(path):
    print(f"Error: File {path} does not exist!")
else:
    doc = docx.Document(path)
    print("--- Paragraphs ---")
    for idx, p in enumerate(doc.paragraphs[:30]):
        txt = p.text.strip()
        if txt:
            print(f"P{idx}: '{txt}'")
            
    print("\n--- Tables ---")
    print(f"Total tables: {len(doc.tables)}")
    for t_idx, t in enumerate(doc.tables):
        first_cell = t.rows[0].cells[0].text.strip().replace('\n', ' ') if len(t.rows) > 0 and len(t.rows[0].cells) > 0 else ""
        print(f"Table {t_idx}: rows={len(t.rows)}, first_cell='{first_cell[:60]}'")
