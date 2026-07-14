import docx

doc = docx.Document("bot/templates/template_contract.docx")

print("=== INSPECTING CONTRACT TEMPLATE PLACEHOLDERS ===")
# Search paragraphs for placeholders
for idx, p in enumerate(doc.paragraphs):
    if "{{" in p.text or "}}" in p.text:
        print(f"P{idx}: '{p.text}'")

# Search table cells (usually the signature block is a table at the end of the document!)
print("\n=== Tables in Contract ===")
for t_idx, t in enumerate(doc.tables):
    print(f"\nTable {t_idx}:")
    for r_idx, row in enumerate(t.rows):
        cells_texts = []
        seen = set()
        for c_idx, cell in enumerate(row.cells):
            if cell._tc not in seen:
                seen.add(cell._tc)
                if "{{" in cell.text or "}}" in cell.text or "Подрядчик" in cell.text:
                    cells_texts.append(f"C{c_idx}: '{cell.text.strip()}'")
        if cells_texts:
            print(f"  Row {r_idx}: {cells_texts}")
