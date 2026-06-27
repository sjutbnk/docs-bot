import docx
import os

replacements = {
    'Кенжаев Бахром Хуррамович': '{{full_name}}',
    'Узбекистана': '{{citizenship}}',
    'Узбекистан': '{{citizenship}}',
    '26.01.1984': '{{birth_date}}',
    'FА': '{{passport_series}}',
    'FA': '{{passport_series}}',
    '1874962': '{{passport_number}}',
    '20.01.2021': '{{passport_issue_date}}',
    'МВД 22232': '{{passport_issued_by}}',
    'Астраханская область, Ахтубинский муниципальный район, городское поселение город Ахтубинск, территория №5, д.2': '{{address}}'
}

def replace_in_paragraph(p):
    text = p.text
    changed = False
    for old, new in replacements.items():
        if old in text:
            text = text.replace(old, new)
            changed = True
    if changed:
        p.text = text

doc = docx.Document('bot/templates/Договор_Кенжаев.docx')
for p in doc.paragraphs:
    replace_in_paragraph(p)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p)

doc.save('bot/templates/template_contract.docx')

try:
    with open('bot/templates/УВЕДОМЛЕНИЕ_О_РАСТОРЖЕНИИ_Кенжаев_Бахром.RTF', 'r', encoding='cp1251') as f:
        rtf_text = f.read()
except UnicodeDecodeError:
    with open('bot/templates/УВЕДОМЛЕНИЕ_О_РАСТОРЖЕНИИ_Кенжаев_Бахром.RTF', 'r', encoding='utf-8', errors='ignore') as f:
        rtf_text = f.read()

for old, new in replacements.items():
    rtf_text = rtf_text.replace(old, new)

with open('bot/templates/template_termination.rtf', 'w', encoding='cp1251', errors='ignore') as f:
    f.write(rtf_text)

print("Templates created successfully.")
