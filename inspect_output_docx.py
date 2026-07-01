import docx
import os

output_dir = 'output/1555793414'
if not os.path.exists(output_dir):
    # try bot/output
    output_dir = 'bot/output/1555793414'

def inspect_output(path):
    print(f"\n=== Inspecting {path} ===")
    if not os.path.exists(path):
        print("File does not exist!")
        return
    doc = docx.Document(path)
    # Print Table 33 Row 1 cells that are not empty
    row = doc.tables[33].rows[1]
    print("Table 33 Row 1:")
    for idx, cell in enumerate(row.cells):
        if cell.text.strip():
            print(f"  Cell {idx}: '{cell.text.strip()}'")
            
    # Also print Table 36 (Patent Validity)
    print("Table 36 Row 0:")
    for idx, cell in enumerate(doc.tables[36].rows[0].cells):
        if cell.text.strip():
            print(f"  Cell {idx}: '{cell.text.strip()}'")

# Find the files
files = [f for f in os.listdir(output_dir) if f.endswith('.docx') and 'Уведомление' in f]
for f in files:
    inspect_output(os.path.join(output_dir, f))
