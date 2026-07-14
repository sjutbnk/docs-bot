import json
import re
import os
import shutil
import tempfile
import time
import docx
from google import genai

import config
import prompts


def _read_docx_as_text(path: str) -> str:
    """Read a .docx file locally using python-docx and return its text content."""
    doc = docx.Document(path)
    lines = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            lines.append(txt)
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            seen = set()
            for cell in row.cells:
                if cell._tc not in seen:
                    seen.add(cell._tc)
                    txt = cell.text.strip().replace('\n', ' ')
                    if txt:
                        row_texts.append(txt)
            if row_texts:
                lines.append(" | ".join(row_texts))
    return "\n".join(lines)


def _parse_json_response(text: str) -> dict:
    """Extract and parse a JSON block from model response text."""
    # Strip markdown code fences with optional language tag and surrounding whitespace
    match = re.search(r'```(?:json)?\s*(.*?)\s*```', text, re.DOTALL)
    if match:
        text = match.group(1)
    data = json.loads(text.strip())
    return _flatten_json(data)


def _flatten_json(data: dict) -> dict:
    """
    Flatten a potentially nested dict returned by the model into a single-level dict.
    The model sometimes wraps fields inside section objects (e.g. "Данные сотрудника": {...}).
    We recursively merge all leaf string values into the top level.
    """
    result = {}
    for key, value in data.items():
        if isinstance(value, dict):
            result.update(_flatten_json(value))
        elif isinstance(value, list):
            config.logger.warning(f"_flatten_json: dropping list value for key '{key}': {value}")
        else:
            result[key] = value
    return result


def _generate_with_fallback(client, contents):
    """
    Generate content using Gemini, falling back to cheaper models on quota/server errors.
    Priority: best quality first, lite models as last resort.
    """
    models = [
        'gemini-2.5-flash',
        'gemini-2.0-flash',
        'gemini-1.5-flash',
        'gemini-2.5-flash-lite',
    ]
    last_exception = None

    for model_name in models:
        for attempt in range(2):
            try:
                config.logger.info(f"Using model: {model_name} (attempt {attempt + 1})")
                response = client.models.generate_content(
                    model=model_name,
                    contents=contents,
                )
                config.logger.info(f"Success with model: {model_name}")
                return response
            except Exception as e:
                last_exception = e
                err_str = str(e)
                config.logger.warning(f"Model {model_name} failed: {err_str}")
                # Retry once on transient errors, then move to next model
                if ("503" in err_str or "429" in err_str) and attempt == 0:
                    time.sleep(3)
                else:
                    break

    raise last_exception


def extract_data_from_images(image_paths: list) -> dict:
    """
    Extract and validate employee/employer data from uploaded document images.

    DOCX files (Partner Cards) are read locally to avoid Gemini mime-type errors.
    Images and PDFs are uploaded to the Gemini Files API.
    A two-phase approach is used: initial extraction then auditor validation.
    """
    if not config.GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY не задан. Получите ключ на aistudio.google.com")

    client = genai.Client(api_key=config.GEMINI_API_KEY)

    uploaded_files = []
    local_text_blocks = []

    for path in image_paths:
        ext = os.path.splitext(path)[1].lower()
        if ext == '.docx':
            try:
                config.logger.info(f"Reading DOCX locally: {path}")
                text = _read_docx_as_text(path)
                local_text_blocks.append(
                    f"=== Содержимое документа {os.path.basename(path)} ===\n{text}"
                )
            except Exception as e:
                config.logger.error(f"Failed to read DOCX {path}: {e}")
        elif ext == '.txt':
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    local_text_blocks.append(
                        f"=== Содержимое документа {os.path.basename(path)} ===\n{f.read()}"
                    )
            except Exception as e:
                config.logger.error(f"Failed to read TXT {path}: {e}")
        else:
            config.logger.info(f"Uploading to Gemini Files API: {path}")
            # Gemini SDK cannot handle non-ASCII characters in file paths.
            # Copy to a temp file with a safe ASCII name before uploading.
            safe_ext = ext if ext else ".bin"
            with tempfile.NamedTemporaryFile(suffix=safe_ext, delete=False) as tmp:
                tmp_path = tmp.name
            try:
                shutil.copy2(path, tmp_path)
                uploaded_files.append(client.files.upload(file=tmp_path))
            finally:
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass

    # Build prompt including any locally-read text
    extraction_prompt = prompts.EXTRACTION_PROMPT
    if local_text_blocks:
        extraction_prompt += "\n\n" + "\n\n".join(local_text_blocks)

    # Phase 1: OCR extraction
    phase1_response = _generate_with_fallback(client, [extraction_prompt] + uploaded_files)
    initial_data = _parse_json_response(phase1_response.text)

    # Phase 2: Auditor cross-check
    validator_prompt = prompts.get_validator_prompt(initial_data)
    phase2_response = _generate_with_fallback(client, [validator_prompt] + uploaded_files)
    validated_data = _parse_json_response(phase2_response.text)

    return validated_data
