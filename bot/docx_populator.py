import docx
import mappings
import utils


# ---------------------------------------------------------------------------
# Low-level cell helpers
# ---------------------------------------------------------------------------

def _set_cell_text(cell, text: str):
    """
    Write text into a table cell while preserving the original run formatting.
    Extra runs beyond the first are blanked (not deleted) to keep formatting intact.
    Extra paragraphs are also blanked rather than removed to avoid layout shifts.
    """
    text = str(text) if text is not None else ""

    # Blank extra paragraphs (keep them for layout)
    for p in cell.paragraphs[1:]:
        for r in p.runs:
            r.text = ""

    if not cell.paragraphs:
        cell.add_paragraph()

    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = text


def _unique_cells(row):
    """Return list of unique cells from a row, skipping merged-cell duplicates."""
    seen = set()
    result = []
    for cell in row.cells:
        if cell._tc not in seen:
            seen.add(cell._tc)
            result.append(cell)
    return result


def _fill_grid(table, row_idx: int, start_col: int, num_cells: int, value: str):
    """Write value character-by-character into consecutive unique grid cells."""
    if not table or row_idx >= len(table.rows):
        return
    cells = _unique_cells(table.rows[row_idx])
    chars = list((value or "").upper())
    for i in range(num_cells):
        col = start_col + i
        if col < len(cells):
            _set_cell_text(cells[col], chars[i] if i < len(chars) else "")


def _fill_date(table, row_idx: int, date_str: str,
               day_idx, month_idx, year_idx, spacer_idx=None):
    """
    Write a DD.MM.YYYY date into separate day/month/year unique grid cells.
    All indices refer to deduplicated cell positions (no merged-cell duplicates).
    """
    if not table or row_idx >= len(table.rows):
        return
    cells = _unique_cells(table.rows[row_idx])

    parts = (date_str or "").split(".")
    day   = parts[0] if len(parts) > 0 else ""
    month = parts[1] if len(parts) > 1 else ""
    year  = parts[2] if len(parts) > 2 else ""

    # Clear all target cells first
    for idx in list(day_idx) + list(month_idx) + list(year_idx) + list(spacer_idx or []):
        if idx < len(cells):
            _set_cell_text(cells[idx], "")

    for i, idx in enumerate(day_idx):
        if idx < len(cells):
            _set_cell_text(cells[idx], day[i] if i < len(day) else "")
    for i, idx in enumerate(month_idx):
        if idx < len(cells):
            _set_cell_text(cells[idx], month[i] if i < len(month) else "")
    for i, idx in enumerate(year_idx):
        if idx < len(cells):
            _set_cell_text(cells[idx], year[i] if i < len(year) else "")


# ---------------------------------------------------------------------------
# Shared employee fields (used by both conclusion & termination notifications)
# ---------------------------------------------------------------------------

def _fill_employee_fields(doc, data: dict):
    """Fill all employee-side fields in the МВД notification template."""
    full_name = str(data.get("full_name") or "")
    name_parts = full_name.split()
    surname    = name_parts[0] if len(name_parts) > 0 else ""
    first_name = name_parts[1] if len(name_parts) > 1 else ""
    patronymic = " ".join(name_parts[2:]) if len(name_parts) > 2 else ""

    # Phone (digits only, up to 11)
    phone = "".join(c for c in str(data.get("phone") or "") if c.isdigit())
    _fill_grid(doc.tables[mappings.PHONE_TABLE], 0, 1, 11, phone)

    # Name fields
    _fill_grid(doc.tables[mappings.SURNAME_TABLE],    0, 1, 28, surname)
    _fill_grid(doc.tables[mappings.NAME_TABLE],       0, 1, 28, first_name)
    _fill_grid(doc.tables[mappings.PATRONYMIC_TABLE], 0, 1, 28, patronymic)
    _fill_grid(doc.tables[mappings.PATRONYMIC_TABLE], 1, 1, 28, "")  # clear line 2

    # Citizenship
    _fill_grid(doc.tables[mappings.CITIZENSHIP_TABLE], 0, 1, 27, data.get("citizenship") or "")

    # Date of birth
    _fill_date(doc.tables[mappings.DOB_TABLE], 0,
               data.get("birth_date") or "", *mappings.DOB_CELLS)

    # Passport series / number / issue date
    _fill_grid(doc.tables[mappings.PASSPORT_TABLE], 0, 1,  7, data.get("passport_series") or "")
    _fill_grid(doc.tables[mappings.PASSPORT_TABLE], 0, 9,  9, data.get("passport_number") or "")
    _fill_date(doc.tables[mappings.PASSPORT_TABLE], 0,
               data.get("passport_issue_date") or "", *mappings.PASSPORT_DATE_CELLS)

    # Passport issued-by (МВД unit code only, split across 3 tables)
    issued = utils.clean_passport_issued_by(data.get("passport_issued_by") or "")
    issued_upper = issued.upper()
    _fill_grid(doc.tables[mappings.PASSPORT_ISSUED_T1], 0, 1, 28, issued_upper[:28])
    _fill_grid(doc.tables[mappings.PASSPORT_ISSUED_T2], 0, 0, 28, issued_upper[28:56])
    _fill_grid(doc.tables[mappings.PASSPORT_ISSUED_T3], 0, 0, 13, issued_upper[56:69])

    # Patent series / number / issue date
    _fill_grid(doc.tables[mappings.PATENT_TABLE], 1, 1,  7, data.get("patent_series") or "")
    _fill_grid(doc.tables[mappings.PATENT_TABLE], 1, 9, 10, data.get("patent_number") or "")
    _fill_date(doc.tables[mappings.PATENT_TABLE], 1,
               data.get("patent_issue_date") or "", *mappings.PATENT_DATE_CELLS)

    # Patent validity period
    _fill_date(doc.tables[mappings.PATENT_VALIDITY_TABLE], 0,
               data.get("patent_issue_date") or "", *mappings.PATENT_VALIDITY_START_CELLS)
    _fill_date(doc.tables[mappings.PATENT_VALIDITY_TABLE], 0,
               data.get("patent_expiry_date") or "", *mappings.PATENT_VALIDITY_END_CELLS)

    # Profession — always cleared (left blank per client request)
    _fill_grid(doc.tables[mappings.PROFESSION_TABLE], 0, 0, 34, "")


def _fill_conclusion_employee_fields(doc, data: dict):
    """
    Populate common employee fields specifically for Conclusion/Termination templates,
    which have a slightly different cell layout compared to the Patent template.
    """
    name_parts = str(data.get("full_name") or "").split()
    surname    = name_parts[0] if len(name_parts) > 0 else ""
    first_name = name_parts[1] if len(name_parts) > 1 else ""
    patronymic = " ".join(name_parts[2:]) if len(name_parts) > 2 else ""

    phone = "".join(c for c in str(data.get("phone") or "") if c.isdigit())
    _fill_grid(doc.tables[21], 0, 1, 11, phone)

    _fill_grid(doc.tables[22], 0, 1, 33, surname)
    _fill_grid(doc.tables[23], 0, 1, 33, first_name)
    _fill_grid(doc.tables[24], 0, 1, 33, patronymic)
    _fill_grid(doc.tables[24], 1, 1, 33, "")

    _fill_grid(doc.tables[25], 0, 1, 33, data.get("citizenship") or "")

    _fill_date(doc.tables[26], 0, data.get("birth_date") or "", *mappings.DOB_CELLS)

    # Conclusion Passport Table is 28
    _fill_grid(doc.tables[28], 0, 1, 7, data.get("passport_series") or "")
    _fill_grid(doc.tables[28], 0, 9, 9, data.get("passport_number") or "")
    _fill_date(doc.tables[28], 0, data.get("passport_issue_date") or "", *mappings.CONCL_PASSPORT_DATE_CELLS)

    issued = utils.clean_passport_issued_by(data.get("passport_issued_by") or "")
    issued_upper = issued.upper()
    _fill_grid(doc.tables[29], 0, 1, 33, issued_upper[:33])
    _fill_grid(doc.tables[30], 0, 0, 33, issued_upper[33:66])
    _fill_grid(doc.tables[31], 0, 0, 33, issued_upper[66:99])

    # Conclusion Patent Table is 33
    _fill_grid(doc.tables[33], 1, 1, 7, data.get("patent_series") or "")
    _fill_grid(doc.tables[33], 1, 9, 10, data.get("patent_number") or "")
    _fill_date(doc.tables[33], 1, data.get("patent_issue_date") or "", *mappings.CONCL_PATENT_DATE_CELLS)

    # Conclusion Patent Validity Table is 36
    _fill_date(doc.tables[36], 0, data.get("patent_issue_date") or "", *mappings.CONCL_PATENT_VALIDITY_START_CELLS)
    _fill_date(doc.tables[36], 0, data.get("patent_expiry_date") or "", *mappings.CONCL_PATENT_VALIDITY_END_CELLS)

    # Profession
    _fill_grid(doc.tables[41], 0, 0, 33, "")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def _fill_conclusion_employer_block(doc, data: dict):
    """Fill employer details for Conclusion and Termination templates (Прил. 7/8)."""
    emp_inn = str(data.get("employer_inn") or "").strip()
    is_default_employer = (emp_inn == "312009347140" or not data.get("employer_name"))

    if not is_default_employer:
        emp_type = str(data.get("employer_type") or "ИП").strip().upper()
        emp_name = str(data.get("employer_name") or "").upper()
        emp_addr = str(data.get("employer_address") or "").upper()

        # Clear existing checkboxes (T4 = Юрлицо, T5 = ИП)
        cells_t4 = _unique_cells(doc.tables[4].rows[0])
        cells_t5 = _unique_cells(doc.tables[5].rows[0])
        _set_cell_text(cells_t4[0], "")
        _set_cell_text(cells_t5[0], "")
        if emp_type == "ИП":
            _set_cell_text(cells_t5[0], "X")
        else:
            _set_cell_text(cells_t4[0], "X")

        # Employer name (T9-T13, 34 chars each)
        for row_offset, t_idx in enumerate(range(9, 14)):
            _fill_grid(doc.tables[t_idx], 0, 0, 34, emp_name[row_offset * 34 : (row_offset + 1) * 34])

        # OGRN (T14)
        _fill_grid(doc.tables[14], 0, 0, 34, str(data.get("employer_ogrn") or ""))

        # Passport (T15, T16)
        if emp_type == "ИП":
            series     = str(data.get("employer_passport_series") or "")
            number     = str(data.get("employer_passport_number") or "")
            issued_by  = str(data.get("employer_passport_issued_by") or "")
            issue_date = str(data.get("employer_passport_issue_date") or "")
            pass_str   = f"ПАСПОРТ {series} {number} {issued_by}".upper()
            _fill_grid(doc.tables[15], 0, 0, 34, pass_str[:34])
            rest = f"{pass_str[34:]} {issue_date}Г.".strip()
            _fill_grid(doc.tables[16], 0, 0, 34, rest[:34])
        else:
            _fill_grid(doc.tables[15], 0, 0, 34, "")
            _fill_grid(doc.tables[16], 0, 0, 34, "")

        # INN (T17)
        _fill_grid(doc.tables[17], 0, 0, 34, str(data.get("employer_inn") or ""))

        # Address (T18-T20)
        _fill_grid(doc.tables[18], 0, 0, 34, emp_addr[:34])
        _fill_grid(doc.tables[19], 0, 0, 34, emp_addr[34:68])
        _fill_grid(doc.tables[20], 0, 0, 34, emp_addr[68:102])

        # Bottom signature (T49)
        cells49 = _unique_cells(doc.tables[49].rows[0])
        if len(cells49) >= 3:
            pure_fio = utils.extract_employer_fio(data.get("employer_name") or "")
            short_name = utils.get_short_name(pure_fio).upper()
            _set_cell_text(cells49[0], emp_type)
            _set_cell_text(cells49[2], short_name)


def fill_conclusion_document(doc, data: dict):
    """Populate МВД notification of contract conclusion."""
    _fill_conclusion_employer_block(doc, data)
    _fill_conclusion_employee_fields(doc, data)
    _fill_date(doc.tables[46], 1,
               "14.05.2026", *mappings.CONCL_CONTRACT_DATE_CELLS)


def fill_termination_document(doc, data: dict):
    """Populate МВД notification of contract termination."""
    _fill_conclusion_employer_block(doc, data)
    _fill_conclusion_employee_fields(doc, data)
    _fill_date(doc.tables[46], 1,
               "30.11.2026", *mappings.CONCL_CONTRACT_DATE_CELLS)


def fill_patent_notification_document(doc, data: dict):
    """
    Populate the patent labour-activity notification (Приказ МВД № 655).
    Employer section is skipped when the default employer (Generalov A.V.) is
    detected — his details are pre-filled in the template to preserve alignment.
    """
    emp_inn = str(data.get("employer_inn") or "").strip()
    is_default_employer = (emp_inn == "312009347140" or not data.get("employer_name"))

    # ── Employer block ──────────────────────────────────────────────────────
    if not is_default_employer:
        emp_type = str(data.get("employer_type") or "ИП").strip().upper()
        emp_name = str(data.get("employer_name") or "").upper()
        emp_addr = str(data.get("employer_address") or "").upper()

        # Name split across 5 rows × 31 chars
        for row_offset, t_idx in enumerate(range(31, 36)):
            _fill_grid(doc.tables[t_idx], 0, 0, 31,
                       emp_name[row_offset * 31 : (row_offset + 1) * 31])

        _fill_grid(doc.tables[36], 0, 0, 31, str(data.get("employer_ogrn") or ""))

        if emp_type == "ИП":
            series     = str(data.get("employer_passport_series") or "")
            number     = str(data.get("employer_passport_number") or "")
            issued_by  = str(data.get("employer_passport_issued_by") or "")
            issue_date = str(data.get("employer_passport_issue_date") or "")
            pass_str   = f"ПАСПОРТ {series} {number} {issued_by}".upper()
            _fill_grid(doc.tables[38], 0, 0, 31, pass_str[:31])
            _fill_grid(doc.tables[38], 1, 0, 31, pass_str[31:62])
            rest = f"{pass_str[62:]} {issue_date}Г.".strip().upper()
            _fill_grid(doc.tables[38], 2, 0, 31, rest[:31])
        else:
            for r in range(4):
                _fill_grid(doc.tables[38], r, 0, 31, "")

        # Employer address (tables 39, 40, 41)
        _fill_grid(doc.tables[39], 0, 0, 31, emp_addr[:31])
        _fill_grid(doc.tables[40], 0, 0, 31, emp_addr[31:62])
        _fill_grid(doc.tables[41], 0, 0, 31, emp_addr[62:93])

        # Employer contact phone from partner card (if available)
        emp_phone = "".join(c for c in str(data.get("employer_phone") or "") if c.isdigit())
        _fill_grid(doc.tables[43], 0, 1, 11, emp_phone)

    # ── Employee block ──────────────────────────────────────────────────────
    full_name  = str(data.get("full_name") or "")
    name_parts = full_name.split()
    surname    = name_parts[0] if len(name_parts) > 0 else ""
    first_name = name_parts[1] if len(name_parts) > 1 else ""
    patronymic = " ".join(name_parts[2:]) if len(name_parts) > 2 else ""

    _fill_grid(doc.tables[3], 0, 1, 28, surname)
    _fill_grid(doc.tables[4], 0, 1, 28, first_name)
    _fill_grid(doc.tables[5], 0, 1, 28, patronymic)
    _fill_grid(doc.tables[5], 1, 1, 28, "")  # clear line 2

    _fill_grid(doc.tables[6], 0, 1, 28, data.get("citizenship") or "")

    # T7: label|D1|D2|sep|M1|M2|sep|Y1|Y2|Y3|Y4  (11 cells, indices 1-10)
    _fill_date(doc.tables[7], 0, data.get("birth_date") or "",
               [1, 2], [4, 5], [7, 8, 9, 10], [3, 6])

    # Passport series (C1-C2=2 chars) / number (C9-C16=8 chars)
    _fill_grid(doc.tables[9], 0, 1, 2, data.get("passport_series") or "")
    _fill_grid(doc.tables[9], 0, 9, 8, data.get("passport_number") or "")

    # T10: label|D1|D2|sep|M1|M2|sep|Y1|Y2|Y3|Y4  (11 cells, same as T7)
    _fill_date(doc.tables[10], 0, data.get("passport_issue_date") or "",
               [1, 2], [4, 5], [7, 8, 9, 10], [3, 6])

    # Passport issued-by: T11 (C1-C25) and T12 (C0-C25)
    issued = utils.clean_passport_issued_by(data.get("passport_issued_by") or "").upper()
    _fill_grid(doc.tables[11], 0, 1, 25, issued[:25])
    _fill_grid(doc.tables[12], 0, 0, 26, issued[25:51])

    # Patent: series (C1-C2=2 chars), number (C9-C18=10 chars),
    # date: C20=D1, C21=D2, C22=sep, C23=M1, C24=M2, C25=sep, C26=Y1, C27=Y2, C28=Y3, C29=Y4
    _fill_grid(doc.tables[13], 0, 1, 2, data.get("patent_series") or "")
    _fill_grid(doc.tables[13], 0, 9, 10, data.get("patent_number") or "")
    _fill_date(doc.tables[13], 0, data.get("patent_issue_date") or "",
               [20, 21], [23, 24], [26, 27, 28, 29], [22, 25])

    # Profession — always blank (3 rows × 31 chars)
    _fill_grid(doc.tables[14], 0, 0, 31, "")
    _fill_grid(doc.tables[15], 0, 0, 31, "")
    _fill_grid(doc.tables[16], 0, 0, 31, "")

    # Place of work (T17 + T18, 31 chars each)
    work_addr = str(data.get("work_address") or data.get("employer_address") or "").upper()
    _fill_grid(doc.tables[17], 0, 0, 31, work_addr[:31])
    _fill_grid(doc.tables[18], 0, 0, 31, work_addr[31:62])

    # Contract type checkbox — cell C2 = "V" for ГПД (civil contract)
    _set_cell_text(doc.tables[20].rows[0].cells[2], "V")

    # Contract date (T22 row 1): label|D1|D2|sep|M1|M2|sep|Y1|Y2|Y3|Y4
    _fill_date(doc.tables[22], 1, "14.05.2026",
               [1, 2], [4, 5], [7, 8, 9, 10], [3, 6])

    # Employee INN (T23, C1-C12)
    _fill_grid(doc.tables[23], 0, 1, 12, str(data.get("inn") or ""))

    # DMS policy series (C1-C3) + number (C7-C18)
    _fill_grid(doc.tables[27], 0, 1, 3, "MRF")
    _fill_grid(doc.tables[27], 0, 7, 12, str(data.get("dms_number") or ""))

    # DMS issue date (T28 row 0, same compact format)
    _fill_date(doc.tables[28], 0, "14.05.2026",
               [1, 2], [4, 5], [7, 8, 9, 10], [3, 6])

    # Employee contact phone (T29, C1-C10 = 10 digits)
    emp_phone = "".join(c for c in str(data.get("phone") or "") if c.isdigit())
    _fill_grid(doc.tables[29], 0, 1, 11, emp_phone)

    # Submitter name (T44) and submission date (T45)
    # T45: «|day|»|month|20|YY|г.  — C1=day(2), C3=month(2), C5=last2digits of year
    _set_cell_text(doc.tables[44].rows[0].cells[1], full_name.upper())
    _set_cell_text(doc.tables[45].rows[0].cells[1], "14")
    _set_cell_text(doc.tables[45].rows[0].cells[3], "05")
    _set_cell_text(doc.tables[45].rows[0].cells[5], "26")
