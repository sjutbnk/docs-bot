import os
import shutil
import tempfile
import docx
from docxtpl import DocxTemplate

import config
import utils
from docx_populator import (
    fill_conclusion_document,
    fill_termination_document,
    fill_patent_notification_document,
)


def generate_documents(data: dict, output_dir: str) -> tuple:
    """
    Generate all four employment documents for a single employee.

    Returns:
        Tuple of (contract_path, conclusion_path, termination_path, patent_notif_path)
    """
    os.makedirs(output_dir, exist_ok=True)

    # 1. Merge default employer if no Partner Card was uploaded
    if not data.get("employer_name") or not str(data.get("employer_name")).strip():
        for key, val in config.DEFAULT_EMPLOYER.items():
            data.setdefault(key, val)

    # 3. Clean passport issued-by to МВД unit code
    if data.get("passport_issued_by"):
        data["passport_issued_by"] = utils.clean_passport_issued_by(
            data["passport_issued_by"]
        )

    # Clean any duplicate prefixes in the employer name that OCR might have merged
    if data.get("employer_name"):
        data["employer_name"] = utils.clean_employer_name(data["employer_name"])

    # 4. Normalize all dates to strict DD.MM.YYYY format (pad missing leading zeros,
    #    handle slash/dash separators). This prevents cells from being written with
    #    single-digit day/month which would visually shift the date in the form.
    for date_key in ("birth_date", "passport_issue_date",
                     "patent_issue_date", "patent_expiry_date",
                     "employer_passport_issue_date"):
        if data.get(date_key):
            data[date_key] = utils.normalize_date(str(data[date_key]))

    # 2. Compute patent dates if one of them is missing or they are identical
    issue_date = str(data.get("patent_issue_date") or "").strip()
    expiry_date = str(data.get("patent_expiry_date") or "").strip()
    if issue_date and (not expiry_date or expiry_date == issue_date):
        data["patent_expiry_date"] = utils.compute_patent_expiry_date(issue_date)
    elif expiry_date and not issue_date:
        data["patent_issue_date"] = utils.compute_patent_issue_date(expiry_date)

    safe_name = (data.get("full_name") or "Сотрудник").replace(" ", "_")

    # ── Contract (Jinja2 template) ──────────────────────────────────────────
    contract_data = data.copy()
    
    # Dates formatting for contract
    c_date = str(data.get("contract_date") or "14.05.2026").strip()
    c_end = str(data.get("contract_end_date") or "30.11.2026").strip()
    
    contract_data["contract_start_date"] = c_date
    contract_data["contract_end_date"]   = c_end
    contract_data["contract_date_ru"]    = utils.format_date_ru(c_date)
    contract_data["contract_start_date_ru"] = utils.format_date_ru(c_date)
    contract_data["contract_end_date_ru"]   = utils.format_date_ru(c_end)
    
    # Profession
    prof = str(data.get("profession") or "Овощевод").strip()
    contract_data["profession_ru"] = prof.capitalize()

    # Bank info from partner card
    contract_data["employer_rs"] = str(data.get("employer_account") or "")
    contract_data["employer_ks"] = str(data.get("employer_corr_account") or "")
    contract_data["employer_bik"] = str(data.get("employer_bik") or "")
    contract_data["employer_bank"] = str(data.get("employer_bank") or "")

    contract_data["short_name"] = utils.get_short_name(data.get("full_name") or "")
    
    # Extract clean FIO for short name (e.g. "Ким В.Р.")
    # employer_name at this point is already cleaned by clean_employer_name,
    # so we just strip the legal prefix and get the FIO.
    pure_fio = utils.extract_employer_fio(data.get("employer_name") or "")
    contract_data["short_employer_name"] = utils.get_short_name(pure_fio)

    # Foreigner address in the contract comes from the Partner Card registration field
    contract_data["address"] = (
        data.get("foreigner_registration_address")
        or data.get("work_address")
        or data.get("employer_address")
        or data.get("address")
        or ""
    )

    tpl_contract = os.path.join(config.TEMPLATES_DIR, "template_contract.docx")
    doc_c = DocxTemplate(tpl_contract)
    doc_c.render(contract_data)
    contract_path = os.path.join(output_dir, f"Договор_прием_{safe_name}.docx")
    doc_c.save(contract_path)

    # ── Conclusion notification (grid-fill) ─────────────────────────────────
    tpl_concl = os.path.join(config.TEMPLATES_DIR, "template_conclusion.docx")
    doc_concl = docx.Document(tpl_concl)
    fill_conclusion_document(doc_concl, data)
    conclusion_path = os.path.join(output_dir, f"Уведомление_прием_{safe_name}.docx")
    doc_concl.save(conclusion_path)

    # ── Termination notification (grid-fill) ────────────────────────────────
    tpl_term = os.path.join(config.TEMPLATES_DIR, "template_termination.docx")
    doc_term = docx.Document(tpl_term)
    fill_termination_document(doc_term, data)
    termination_path = os.path.join(output_dir, f"Уведомление_расторжение_{safe_name}.docx")
    doc_term.save(termination_path)

    # ── Patent notification (grid-fill) ─────────────────────────────────────
    tpl_patent = os.path.join(config.TEMPLATES_DIR, "template_patent_notification.docx")
    doc_pn = docx.Document(tpl_patent)
    fill_patent_notification_document(doc_pn, data)
    
    citizen = str(data.get("citizenship") or "").strip().lower()
    if "узбек" in citizen:
        cit_str = "от_узбека"
    elif "таджик" in citizen:
        cit_str = "от_таджика"
    elif citizen:
        cit_str = f"от_{citizen}"
    else:
        cit_str = "от_иностранца"
        
    patent_path = os.path.join(output_dir, f"Уведомление_{cit_str}_{safe_name}.docx")
    doc_pn.save(patent_path)

    return contract_path, conclusion_path, termination_path, patent_path
