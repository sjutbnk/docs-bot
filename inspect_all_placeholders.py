import docx
import re

doc = docx.Document('bot/templates/template_contract.docx')
placeholders = set()

# Find placeholders in paragraphs
for p in doc.paragraphs:
    matches = re.findall(r'\{\{(.*?)\}\}', p.text)
    for m in matches:
        placeholders.add(m.strip())
        
# Find placeholders in tables
for t in doc.tables:
    for r in t.rows:
        for cell in r.cells:
            matches = re.findall(r'\{\{(.*?)\}\}', cell.text)
            for m in matches:
                placeholders.add(m.strip())
                
print("Jinja placeholders in template_contract.docx:")
for p in sorted(placeholders):
    print(f"  {p}")
